from io import BytesIO
def parse_bytes(filename: str, data: bytes) -> str:
    ext = filename.rsplit('.', 1)[-1].lower()

    if ext == "docx":
        from docx import Document
        doc = Document(BytesIO(data))
        parts = []
        parts.extend(p.text for p in doc.paragraphs if p.text)
        for t in doc.tables:
            for row in t.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)

    if ext in ("txt", "md"):
        for enc in ("utf-8", "utf-8-sig", "cp1250", "iso-8859-2", "utf-16"):
            try:
                return data.decode(enc)
            except UnicodeDecodeError:
                continue
        return data.decode("utf-8", errors="ignore")

    if ext == "pdf":
        from pdfminer.high_level import extract_text
        return extract_text(BytesIO(data)) or ""

    if ext == "csv":
        import pandas as pd
        from io import StringIO
        df = pd.read_csv(StringIO(data.decode("utf-8", errors="ignore")))
        return "\n".join(" | ".join(map(str, row)) for row in df.values)

    return ""